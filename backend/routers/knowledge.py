from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, distinct
from database import get_db
from models import KnowledgeBase, KnowledgeEntry, DocumentUpload, User, SystemSettings
from schemas import KBCreate, KBUpdate, KBOut, EntryCreate, EntryUpdate, EntryOut
from auth import get_current_user, require_admin
from typing import List, Optional
import json
from io import BytesIO

router = APIRouter(prefix="/api/admin/kb", tags=["知识库管理"])


async def _fill_kb_stats(kb: KnowledgeBase, db: AsyncSession):
    entry_count_result = await db.execute(
        select(func.count()).where(KnowledgeEntry.kb_id == kb.id)
    )
    kb.entry_count = entry_count_result.scalar() or 0
    doc_count_result = await db.execute(
        select(func.count(distinct(KnowledgeEntry.source_doc))).where(KnowledgeEntry.kb_id == kb.id)
    )
    kb.doc_count = doc_count_result.scalar() or 0


@router.get("/qa-template")
async def download_qa_template(
    current_user: User = Depends(get_current_user),
):
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    title = doc.add_heading('问答对参考模板', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph('')
    note = doc.add_paragraph()
    note_run = note.add_run('说明：请按照以下格式填写问答对，系统会自动识别并解析。支持以下问答标记格式：')
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    formats = [
        '格式一：Q: 问题 / A: 答案',
        '格式二：问：问题 / 答：答案',
        '格式三：问题：问题内容 / 答案：答案内容',
    ]
    for fmt in formats:
        p = doc.add_paragraph(fmt, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    doc.add_heading('示例一：产品常见问题', level=2)

    examples = [
        ("Q: KnowQ是什么？", "A: KnowQ智答星是一款AI知识库问答平台，支持创建多个知识库，配置大模型与搜索策略，并一键生成面向最终用户的AI问答前端。"),
        ("Q: 如何创建知识库？", "A: 登录管理后台，在左侧导航栏点击「知识库管理」，然后点击「新建知识库」按钮，填写知识库名称和描述即可创建。"),
        ("Q: 支持哪些文档格式？", "A: 目前支持TXT、PDF、Word(.docx)、Markdown等常见文档格式。系统会自动解析文档内容并分块入库。"),
    ]
    for q, a in examples:
        p = doc.add_paragraph()
        q_run = p.add_run(q)
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        p2 = doc.add_paragraph()
        a_run = p2.add_run(a)
        a_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph('')
    doc.add_heading('示例二：公司规章制度', level=2)

    examples2 = [
        ("问：公司上班时间是什么？", "答：公司标准上班时间为周一至周五 9:00-18:00，午休时间 12:00-13:00。弹性工作制下可在8:30-10:00之间打卡上班。"),
        ("问：如何申请年假？", "答：员工入职满一年后享有5天年假，满五年10天，满十年15天。申请年假需提前3个工作日在OA系统提交申请，经直属领导审批后生效。"),
        ("问：报销流程是什么？", "答：1. 收集相关票据和凭证；2. 在OA系统填写报销单，上传凭证照片；3. 提交至部门经理审批；4. 审批通过后转财务处理；5. 财务审核通过后5个工作日内打款。"),
    ]
    for q, a in examples2:
        p = doc.add_paragraph()
        q_run = p.add_run(q)
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        p2 = doc.add_paragraph()
        a_run = p2.add_run(a)
        a_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph('')
    doc.add_heading('请在下方填写您的问答对', level=2)

    for _ in range(3):
        doc.add_paragraph('')
        p = doc.add_paragraph()
        q_run = p.add_run('Q: ')
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        p.add_run('（请在此处填写问题）')

        p2 = doc.add_paragraph()
        a_run = p2.add_run('A: ')
        a_run.bold = True
        a_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        p2.add_run('（请在此处填写答案）')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=qa_template.docx"},
    )


@router.get("", response_model=List[KBOut])
async def list_kbs(
    keyword: str = Query(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    q = select(KnowledgeBase).order_by(KnowledgeBase.created_at.desc())
    if keyword:
        q = q.where(KnowledgeBase.name.contains(keyword))
    result = await db.execute(q)
    kbs = result.scalars().all()
    out = []
    for kb in kbs:
        await _fill_kb_stats(kb, db)
        out.append(KBOut.model_validate(kb))
    return out


@router.post("", response_model=KBOut)
async def create_kb(
    req: KBCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    kb = KnowledgeBase(
        name=req.name,
        description=req.description,
        created_by=current_user.id,
    )
    db.add(kb)
    await db.commit()
    await db.refresh(kb)
    return KBOut.model_validate(kb)


@router.get("/{kb_id}", response_model=KBOut)
async def get_kb(
    kb_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    await _fill_kb_stats(kb, db)
    return KBOut.model_validate(kb)


@router.put("/{kb_id}", response_model=KBOut)
async def update_kb(
    kb_id: str,
    req: KBUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    if req.name is not None:
        kb.name = req.name
    if req.description is not None:
        kb.description = req.description
    if req.status is not None:
        kb.status = req.status
    await db.commit()
    await db.refresh(kb)
    return KBOut.model_validate(kb)


@router.delete("/{kb_id}")
async def delete_kb(
    kb_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")
    await db.delete(kb)
    await db.commit()
    return {"message": "删除成功"}


@router.get("/{kb_id}/entries", response_model=List[EntryOut])
async def list_entries(
    kb_id: str,
    keyword: str = Query(None),
    entry_type: str = Query(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    q = select(KnowledgeEntry).where(KnowledgeEntry.kb_id == kb_id).order_by(KnowledgeEntry.created_at.desc())
    if keyword:
        q = q.where(
            (KnowledgeEntry.question.contains(keyword)) | (KnowledgeEntry.content.contains(keyword))
        )
    if entry_type:
        q = q.where(KnowledgeEntry.type == entry_type)
    result = await db.execute(q)
    entries = result.scalars().all()
    return [EntryOut.model_validate(e) for e in entries]


@router.post("/{kb_id}/entries", response_model=EntryOut)
async def create_entry(
    kb_id: str,
    req: EntryCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    if not kb_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="知识库不存在")
    entry = KnowledgeEntry(
        kb_id=kb_id,
        type=req.type,
        question=req.question,
        content=req.content,
        source_doc=req.source_doc,
    )
    db.add(entry)
    await db.commit()
    await db.refresh(entry)
    return EntryOut.model_validate(entry)


@router.put("/{kb_id}/entries/{entry_id}", response_model=EntryOut)
async def update_entry(
    kb_id: str,
    entry_id: str,
    req: EntryUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id, KnowledgeEntry.kb_id == kb_id)
    )
    entry = result.scalar_one_or_none()
    if not entry:
        raise HTTPException(status_code=404, detail="条目不存在")
    if req.question is not None:
        entry.question = req.question
    if req.content is not None:
        entry.content = req.content
    if req.source_doc is not None:
        entry.source_doc = req.source_doc
    await db.commit()
    await db.refresh(entry)
    return EntryOut.model_validate(entry)


@router.delete("/{kb_id}/entries/{entry_id}")
async def delete_entry(
    kb_id: str,
    entry_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id, KnowledgeEntry.kb_id == kb_id)
    )
    entry = result.scalar_one_or_none()
    if not entry:
        raise HTTPException(status_code=404, detail="条目不存在")
    await db.delete(entry)
    await db.commit()
    return {"message": "删除成功"}


@router.post("/{kb_id}/import")
async def import_document(
    kb_id: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = kb_result.scalar_one_or_none()
    if not kb:
        raise HTTPException(status_code=404, detail="知识库不存在")

    from services.document_parser import parse_document
    from services.qa_engine import get_system_settings
    content = await file.read()
    filename = file.filename or "unknown.txt"

    settings = await get_system_settings(db)

    try:
        entries = parse_document(
            filename, content,
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            qa_detect_lines=settings.qa_detect_lines,
            qa_min_markers=settings.qa_min_markers,
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文档解析失败: {str(e)}")

    chunk_count = 0
    for entry_data in entries:
        entry = KnowledgeEntry(
            kb_id=kb_id,
            type=entry_data.get("type", "chunk"),
            question=entry_data.get("question"),
            content=entry_data["content"],
            source_doc=filename,
            chunk_index=chunk_count if entry_data.get("type") == "chunk" else None,
        )
        db.add(entry)
        chunk_count += 1

    await db.commit()

    return {"message": f"成功导入 {chunk_count} 条知识条目", "count": chunk_count}
